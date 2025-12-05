"""Resume parser module for extracting text from PDF, DOCX, and TXT files."""

import pdfplumber
import re
from typing import Optional
from pathlib import Path
import logging
from docx import Document

from ..models import ResumeData


logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse PDF resumes and extract structured information."""
    
    def __init__(self):
        """Initialize the resume parser."""
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
    
    def parse_pdf(self, pdf_path: str) -> ResumeData:
        """
        Parse a PDF resume and extract text content.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            ResumeData object with extracted information
            
        Raises:
            FileNotFoundError: If the PDF file doesn't exist
            Exception: If PDF parsing fails
        """
        path = Path(pdf_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {pdf_path}")
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # Extract text from all pages
                raw_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        raw_text += page_text + "\n"
                
                if not raw_text.strip():
                    raise ValueError("No text could be extracted from the PDF")
                
                # Extract structured information
                name = self._extract_name(raw_text)
                email = self._extract_email(raw_text)
                skills = self._extract_skills(raw_text)
                
                logger.info(f"Successfully parsed resume: {pdf_path}")
                
                return ResumeData(
                    raw_text=raw_text,
                    name=name,
                    email=email,
                    skills=skills
                )
                
        except Exception as e:
            logger.error(f"Error parsing PDF {pdf_path}: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_pdf_bytes(self, pdf_bytes: bytes) -> ResumeData:
        """
        Parse a PDF resume from bytes.
        
        Args:
            pdf_bytes: PDF file content as bytes
            
        Returns:
            ResumeData object with extracted information
        """
        try:
            import io
            pdf_file = io.BytesIO(pdf_bytes)
            
            with pdfplumber.open(pdf_file) as pdf:
                raw_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        raw_text += page_text + "\n"
                
                if not raw_text.strip():
                    raise ValueError("No text could be extracted from the PDF")
                
                name = self._extract_name(raw_text)
                email = self._extract_email(raw_text)
                skills = self._extract_skills(raw_text)
                
                logger.info("Successfully parsed resume from bytes")
                
                return ResumeData(
                    raw_text=raw_text,
                    name=name,
                    email=email,
                    skills=skills
                )
                
        except Exception as e:
            logger.error(f"Error parsing PDF bytes: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_docx(self, docx_path: str) -> ResumeData:
        """
        Parse a DOCX (Word) resume and extract text content.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            ResumeData object with extracted information
            
        Raises:
            FileNotFoundError: If the DOCX file doesn't exist
            Exception: If DOCX parsing fails
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            # Extract text from all paragraphs
            raw_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    raw_text += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            raw_text += cell.text + "\n"
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            # Extract structured information
            name = self._extract_name(raw_text)
            email = self._extract_email(raw_text)
            skills = self._extract_skills(raw_text)
            
            logger.info(f"Successfully parsed resume: {docx_path}")
            
            return ResumeData(
                raw_text=raw_text,
                name=name,
                email=email,
                skills=skills
            )
                
        except Exception as e:
            logger.error(f"Error parsing DOCX {docx_path}: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_docx_bytes(self, docx_bytes: bytes) -> ResumeData:
        """
        Parse a DOCX resume from bytes.
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            ResumeData object with extracted information
        """
        try:
            import io
            docx_file = io.BytesIO(docx_bytes)
            
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            raw_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    raw_text += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            raw_text += cell.text + "\n"
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            name = self._extract_name(raw_text)
            email = self._extract_email(raw_text)
            skills = self._extract_skills(raw_text)
            
            logger.info("Successfully parsed resume from DOCX bytes")
            
            return ResumeData(
                raw_text=raw_text,
                name=name,
                email=email,
                skills=skills
            )
                
        except Exception as e:
            logger.error(f"Error parsing DOCX bytes: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_txt(self, txt_path: str) -> ResumeData:
        """
        Parse a TXT resume and extract text content.
        
        Args:
            txt_path: Path to the TXT file
            
        Returns:
            ResumeData object with extracted information
            
        Raises:
            FileNotFoundError: If the TXT file doesn't exist
            Exception: If TXT parsing fails
        """
        path = Path(txt_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {txt_path}")
        
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                raw_text = f.read()
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the TXT file")
            
            # Extract structured information
            name = self._extract_name(raw_text)
            email = self._extract_email(raw_text)
            skills = self._extract_skills(raw_text)
            
            logger.info(f"Successfully parsed resume: {txt_path}")
            
            return ResumeData(
                raw_text=raw_text,
                name=name,
                email=email,
                skills=skills
            )
                
        except Exception as e:
            logger.error(f"Error parsing TXT {txt_path}: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_txt_bytes(self, txt_bytes: bytes) -> ResumeData:
        """
        Parse a TXT resume from bytes.
        
        Args:
            txt_bytes: TXT file content as bytes
            
        Returns:
            ResumeData object with extracted information
        """
        try:
            raw_text = txt_bytes.decode('utf-8')
            
            if not raw_text.strip():
                raise ValueError("No text could be extracted from the TXT file")
            
            name = self._extract_name(raw_text)
            email = self._extract_email(raw_text)
            skills = self._extract_skills(raw_text)
            
            logger.info("Successfully parsed resume from TXT bytes")
            
            return ResumeData(
                raw_text=raw_text,
                name=name,
                email=email,
                skills=skills
            )
                
        except Exception as e:
            logger.error(f"Error parsing TXT bytes: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def parse_resume_bytes(self, file_bytes: bytes, file_name: str) -> ResumeData:
        """
        Parse a resume from bytes, automatically detecting the file format.
        
        Args:
            file_bytes: File content as bytes
            file_name: File name with extension (used to determine format)
            
        Returns:
            ResumeData object with extracted information
            
        Raises:
            ValueError: If file format is not supported
            Exception: If parsing fails
        """
        file_lower = file_name.lower()
        
        if file_lower.endswith('.pdf'):
            return self.parse_pdf_bytes(file_bytes)
        elif file_lower.endswith('.docx'):
            return self.parse_docx_bytes(file_bytes)
        elif file_lower.endswith('.txt'):
            return self.parse_txt_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {file_name}. Supported formats: PDF, DOCX, TXT")
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        match = self.email_pattern.search(text)
        return match.group(0) if match else None
    
    def _extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name from resume text.
        Assumes name is in the first few lines.
        """
        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            # Simple heuristic: name is usually short and at the top
            if line and len(line.split()) <= 4 and len(line) < 50:
                # Avoid lines with common resume keywords
                if not any(keyword in line.lower() for keyword in 
                          ['resume', 'cv', 'curriculum', 'email', 'phone', 'address']):
                    return line
        return None
    
    def _extract_skills(self, text: str) -> list:
        """
        Extract technical skills from resume text.
        This is a simple keyword-based approach.
        """
        # Common technical skills and technologies
        skill_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'fastapi',
            'sql', 'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
            'git', 'jenkins', 'ci/cd', 'agile', 'scrum',
            'machine learning', 'deep learning', 'ai', 'nlp', 'computer vision',
            'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
            'rest api', 'graphql', 'microservices', 'linux', 'bash'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill)
        
        return list(set(found_skills))  # Remove duplicates


class ResumeParserError(Exception):
    """Custom exception for resume parsing errors."""
    pass
