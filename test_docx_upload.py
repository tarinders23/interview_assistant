#!/usr/bin/env python3
"""
Test script to verify DOCX and TXT file upload and parsing through the API.
"""
import requests
import tempfile
import os
from docx import Document
from pathlib import Path

# Configuration
API_URL = "http://localhost:8000"
ENDPOINT = "/api/v1/generate-questions"

def create_test_docx():
    """Create a test DOCX resume file."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False) as tmp:
        doc = Document()
        doc.add_paragraph('Jane Developer', style='Heading 1')
        doc.add_paragraph('jane.developer@example.com')
        doc.add_paragraph('Phone: +1-555-123-4567')
        doc.add_paragraph('')
        doc.add_paragraph('SKILLS')
        doc.add_paragraph('Python, JavaScript, React, FastAPI, PostgreSQL, Docker')
        doc.add_paragraph('')
        doc.add_paragraph('EXPERIENCE')
        doc.add_paragraph('Senior Full-Stack Developer at TechCorp (2020-2024)')
        doc.add_paragraph('Built scalable web applications using React and FastAPI')
        doc.add_paragraph('')
        doc.add_paragraph('Full-Stack Developer at StartupXYZ (2018-2020)')
        doc.add_paragraph('Developed microservices architecture using Python and Node.js')
        
        tmp_path = tmp.name
        doc.save(tmp_path)
        return tmp_path

def create_test_txt():
    """Create a test TXT resume file."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as tmp:
        tmp.write("""John Developer
john.dev@example.com
+1-555-987-6543

TECHNICAL SKILLS
JavaScript, TypeScript, Node.js, Express, React, Python, Django, PostgreSQL

PROFESSIONAL EXPERIENCE
Senior Backend Engineer at CloudCorp (2021-2024)
- Developed RESTful APIs serving 100k+ daily users
- Optimized database queries reducing latency by 40%

Backend Developer at WebServices Inc (2019-2021)
- Built microservices using Node.js and Docker
- Implemented CI/CD pipelines with GitHub Actions
""")
        return tmp.name

def test_file_upload(file_path, file_type):
    """Test uploading a file through the API."""
    print(f"\n{'='*60}")
    print(f"Testing {file_type.upper()} Upload")
    print(f"{'='*60}")
    
    file_name = os.path.basename(file_path)
    print(f"File: {file_name}")
    print(f"Size: {os.path.getsize(file_path)} bytes")
    
    with open(file_path, 'rb') as f:
        files = {'resume': (file_name, f)}
        data = {
            'job_description': 'We are looking for a senior full-stack developer with 5+ years of experience in building scalable web applications using modern tech stack.',
            'round_type': 'technical',
            'difficulty': 'advanced',
            'num_questions': 3,
            'api_key': 'test-key'  # This will fail but we just want to test the file parsing
        }
        
        try:
            response = requests.post(f"{API_URL}{ENDPOINT}", files=files, data=data, timeout=10)
            print(f"Status Code: {response.status_code}")
            print(f"Response Headers: {dict(response.headers)}")
            
            try:
                response_json = response.json()
                print(f"Response JSON: {response_json}")
                
                if 'detail' in response_json:
                    print(f"\n✗ ERROR: {response_json['detail']}")
                    return False
                else:
                    print(f"\n✓ File processing succeeded!")
                    return True
            except:
                print(f"Response Text: {response.text[:200]}")
                return False
                
        except Exception as e:
            print(f"✗ Request failed: {e}")
            return False

def main():
    """Run all tests."""
    print("DOCX and TXT File Upload Test Suite")
    print("====================================")
    print(f"Target API: {API_URL}{ENDPOINT}")
    
    # Create test files
    print("\nCreating test files...")
    docx_file = create_test_docx()
    txt_file = create_test_txt()
    
    print(f"✓ Created test DOCX: {docx_file}")
    print(f"✓ Created test TXT: {txt_file}")
    
    # Test uploads
    results = {}
    results['DOCX'] = test_file_upload(docx_file, 'docx')
    results['TXT'] = test_file_upload(txt_file, 'txt')
    
    # Cleanup
    print(f"\n\nCleaning up...")
    os.unlink(docx_file)
    os.unlink(txt_file)
    print("✓ Cleaned up test files")
    
    # Summary
    print(f"\n{'='*60}")
    print("TEST SUMMARY")
    print(f"{'='*60}")
    for file_type, passed in results.items():
        status = "✓ PASS" if passed else "✗ FAIL"
        print(f"{file_type:10} {status}")
    
    all_passed = all(results.values())
    if all_passed:
        print("\n✓ All tests passed!")
    else:
        print("\n✗ Some tests failed. Check the output above for details.")
    
    return all_passed

if __name__ == '__main__':
    main()
